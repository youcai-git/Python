import argparse
import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from future.moves import sys
from openpyxl import load_workbook
from docx.shared import RGBColor


def read_de_info(excel_path):
    ws = load_workbook(excel_path).worksheets[0]  # 读取数据元sheet页，默认第一个sheet页
    de_info_dic = {}  # 存储数据元属性信息
    de_sort__ = ws.cell(1, 1).value  # 数据元分类：人员类
    de_name__ = ws.cell(1, 2).value  # 中文名称：公民身份号码
    de_quote_ = ws.cell(1, 3).value  # 引用数据元规范：政务数据元
    de_i_mark = ws.cell(1, 4).value  # 内部标识符：010001
    de_mark__ = ws.cell(1, 5).value  # 标识符：GMSFHM
    de_def___ = ws.cell(1, 6).value  # 定义：赋码机关为每个公民给出的唯一的、终身不变的法定标识号码
    de_sync__ = ws.cell(1, 7).value  # 同义名称：身份证号、居民身份证号
    de_type__ = ws.cell(1, 8).value  # 数据类型：字符型
    de_format = ws.cell(1, 9).value  # 格式：c18
    de_range_ = ws.cell(1, 10).value  # 值域：符合GB 11643《公民身份号码》
    de_unit__ = ws.cell(1, 11).value  # 计量单位：
    de_remark = ws.cell(1, 12).value  # 备注：

    if de_sort__ != '数据元分类' or de_name__ != '中文名称' or de_quote_ != '引用数据元规范' \
            or de_i_mark != '内部标识符' or de_mark__ != '拼音/标识符' or de_def___ != '定义' \
            or de_sync__ != '同义名称' or de_type__ != '数据类型' or de_format != '格式' \
            or de_range_ != '值域' or de_unit__ != '计量单位' or de_remark != '备注':
        print("Error：Excel文件表头格式不对")
        return 0
    for row in range(2, ws.max_row + 1):
        de_sort__ = str(ws.cell(row, 1).value).replace("\n", "").replace("\r", "")  # 数据元分类
        de_name__ = str(ws.cell(row, 2).value).replace("\n", "").replace("\r", "")  # 数据元中文名称
        de_quote_ = str(ws.cell(row, 3).value).replace("\n", "").replace("\r", "")  # 引用数据元规范
        de_i_mark = str(ws.cell(row, 4).value).replace("\n", "").replace("\r", "")  # 内部标识符
        de_mark__ = str(ws.cell(row, 5).value).replace("\n", "").replace("\r", "")  # 标识符
        de_def___ = str(ws.cell(row, 6).value).replace("\n", "").replace("\r", "")  # 数据元定义
        de_sync__ = str(ws.cell(row, 7).value).replace("\n", "").replace("\r", "")  # 数据元类型
        de_type__ = str(ws.cell(row, 8).value).replace("\n", "").replace("\r", "")  # 数据元格式
        de_format = str(ws.cell(row, 9).value).replace("\n", "").replace("\r", "")  # 数据元格式
        de_range_ = str(ws.cell(row, 10).value).replace("\n", "").replace("\r", "")  # 值域
        de_unit__ = str(ws.cell(row, 11).value).replace("\n", "").replace("\r", "")  # 计量单位
        de_remark = str(ws.cell(row, 12).value).replace("\n", "").replace("\r", "")  # 备注

        key = int(de_i_mark)  # 选择内部标示符作为key
        if key not in de_info_dic:
            de_info_dic.setdefault(key, ['中文名称：' + de_name__, '内部标识符：' + de_i_mark,
                                         '标识符：' + de_mark__, '定义：' + de_def___,
                                         '同义名称：' + de_sync__, '数据类型：' + de_type__,
                                         '格式：' + de_format, '值域：' + de_range_,
                                         '计量单位：' + de_unit__, '备注：' + de_remark])
    de_info_dic = dict(sorted(de_info_dic.items(), key=lambda x: x[0]))  # 以内部标识符升序排序
    return de_info_dic


def create_de_standard(de_info_dic, de_classify, de_number):
    doc = Document()
    # 文档全局字体设置
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal'].font.size = Pt(10.5)  # 五号字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    title_index = 1  # 标题编号，从 1 开始
    for key in de_info_dic:
        if str(key).startswith(de_classify):
            title_content = de_info_dic[key][0].replace("中文名称：", "")   # title名称
            print('cvt start: ' + de_number + '.' + str(title_index) + ' ' + title_content)
            run = doc.add_heading('', level=4).add_run(
                de_number + '.' + str(title_index) + ' ' + title_content)
            run.font.name = '黑体'  # 设置字体格式.注：这个好像设置 run 中的西文字体
            run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # 设置中文字体
            run.font.size = Pt(14)  # 设置字体大小事故数据元V3.0.xlsx
            run.font.bold = False  # 设置加粗
            run.font.italic = False  # 设置斜体
            run.font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色;需导入 rgb 颜色模块
            run.font.underline = False  # 设置下划线
            run.font.outline = False  # 设置轮廓线
            run.font.shadow = False  # 设置阴影
            run.font.strike = False  # 删除线
            run.font.double_strike = False  # 双删除线
            run.font.subscript = False  # 设置下标
            run.font.superscript = False  # 设置上标

            title_index += 1
            # 数据元具体内容
            for content in de_info_dic[key]:
                par = doc.add_paragraph(text=content.replace("None", ""))  # None值用''代替
                par.paragraph_format.line_spacing = 1.5
                par.paragraph_format.space_before = Pt(0)  # 段前间距
                par.paragraph_format.space_after = Pt(0)  # 段后间距
                par.paragraph_format.left_indent = Pt(21)  # 左缩进2个字符
                par.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐

    doc.save("de_output.docx")


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='将数据元excel形式转换成标准word格式')
    parser.add_argument('--excel_path', '-f', type=str, default='事故数据元V3.0.xlsx', help='指定excel输入路径')
    parser.add_argument('--de_classify', '-c', type=str, default='', help='数据元分类--内部标识符')
    parser.add_argument('--de_number', '-n', type=str, default='5.4', help='数据元标题编号')
    args = parser.parse_args()
    print("cvt starting....\n")
    create_de_standard(read_de_info(args.excel_path), args.de_classify, args.de_number)
    print("cvt end....")
    print('数据元标准word格式：de_output.docx...生成成功')
    print('输出word所在目录为：' + os.path.abspath(sys.argv[0]))
